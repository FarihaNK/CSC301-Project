<<<<<<< HEAD
# import argparse
# # from langchain_chroma import Chroma
# # from langchain.vectorstores import Chroma
# from langchain_chroma import Chroma

# from langchain.prompts import ChatPromptTemplate
# #from langchain_community.llms.ollama import Ollama
# #from langchain.llms import Ollama
# from langchain_ollama import OllamaLLM



# from get_embedding import get_embedding_function
# from langchain.schema.document import Document
# from docx import Document as DocxDocument
# import os

# CHROMA_PATH = "chroma"

# PROMPT_TEMPLATE = """
# You are a knowledgeable assistant. Use the provided context to answer the question thoroughly and accurately. Make sure to base your answer only on the information in the context, and clearly reference the specific document Name and page Number when available. When asked for general information, reference all the data that has been provided. If user query information is not present, Let the User know right away.

# Context:
# {context}

# --- 

# Based on the above context, answer the question: {question}

# Please include the source references in your response.
# """


# def main():
#     parser = argparse.ArgumentParser()
#     parser.add_argument("query_text", type=str, nargs='?', default=None, help="The query text.")
#     args = parser.parse_args()
#     if args.query_text:
#         # If a query is provided as an argument, process it once
#         query_rag(args.query_text)
#     else:
#         # Otherwise, enter live mode
#         live_mode()


# def live_mode():
#     print("Entering live mode. Type 'exit' or 'quit' to stop.")
#     while True:
#         query_text = input("Enter your question: ")
#         if query_text.lower() in ['exit', 'quit']:  # ensures all lowercase
#             print("Exiting...")
#             break
#         query_rag(query_text)

# """
# def read_text_file(file_path):
#     try:
#         with open(file_path, 'r', encoding='utf-8') as file:
#             return file.read()
#     except UnicodeDecodeError:
#         with open(file_path, 'r', encoding='iso-8859-1') as file:
#             return file.read()


# def load_txt_files(data_path):
#     txt_documents = []
#     for filename in os.listdir(data_path):
#         if filename.endswith(".txt"):
#             content = read_text_file(os.path.join(data_path, filename))
#             doc_id = f"{filename}"
#             txt_documents.append(Document(page_content=content, metadata={"source": doc_id}))
#     return txt_documents


# def load_docx_files(data_path):
#     docx_documents = []
#     for filename in os.listdir(data_path):
#         if filename.endswith(".docx"):
#             doc = DocxDocument(os.path.join(data_path, filename))
#             content = "\n".join([para.text for para in doc.paragraphs])
#             doc_id = f"{filename}"
#             docx_documents.append(Document(page_content=content, metadata={"source": doc_id}))
#     return docx_documents
#     """


# def query_rag(query_text: str):
#     try:
#         # Prepare the DB.
#         embedding_function = get_embedding_function()
#         db = Chroma(persist_directory=CHROMA_PATH, embedding_function=embedding_function)

#         # Search the DB.
#         results = db.similarity_search_with_score(query_text, k=5)
#         #checks 5 most similar chunks to answer data .

#         # Create context text with source labels
#         context_text = "\n\n---\n\n".join(
#             [f"Source {i + 1}: {doc.page_content}" for i, (doc, _score) in enumerate(results)]
#         )
#         prompt_template = ChatPromptTemplate.from_template(PROMPT_TEMPLATE) # pass tempelate over
#         prompt = prompt_template.format(context=context_text, question=query_text) #langchain function 

#         # Use the model to generate a response
#         #model = Ollama(model="mistral")
#         model = OllamaLLM(model="Mistral")  # 

#         response_text = model.invoke(prompt)

#         # Extract and format the source identifiers
#         sources = []
#         for i, (doc, _score) in enumerate(results):
#             # doc.metadata.get(key, default_value)
#             # if "source" doesn't exist, it defaults to e.g. "Source 1", "Source 2", etc.
#             source_value = doc.metadata.get("source", f"Source {i + 1}")
#             if source_value not in sources:
#                 sources.append(source_value)

#         formatted_sources = "\n".join(sources)
#         formatted_response = f"RESPONSE: {response_text}\nSources used:\n{formatted_sources}"
#         print("")
#         print("")
#         print(formatted_response)
#     except Exception as e:
#         print(f"An error occurred: {e}")


# if __name__ == "__main__":
#     main()


#VERSION 2: ATTEMPTING TO FIX RESPONSE THING: 


import argparse
=======
#STABLE VERSION DATE: THURSDAY MARCH 27TH 11:22 PM 
import argparse
import os
import time
>>>>>>> attempted-AI-integration
from langchain_chroma import Chroma
from langchain.prompts import ChatPromptTemplate
from langchain_ollama import OllamaLLM
from get_embedding import get_embedding_function
from langchain.schema.document import Document
<<<<<<< HEAD
from docx import Document as DocxDocument
import os
=======
from agents import escalate_to_sme
>>>>>>> attempted-AI-integration

CHROMA_PATH = "chroma"

PROMPT_TEMPLATE = """
<<<<<<< HEAD
You are a knowledgeable assistant. Use the provided context to answer the question thoroughly and accurately. Make sure to base your answer only on the information in the context, and clearly reference the specific document Name and page Number when available. When asked for general information, reference all the data that has been provided. If user query information is not present, Let the User know right away.
=======
You are a friendly and helpful assistant with a knack for synthesizing information. Use the context provided below to answer the question in a natural, conversational manner. 
Please do not simply copy the context verbatim; instead, interpret and summarize the key points.

Example:
Context: "The application experienced a brief downtime due to unexpected high traffic and network issues."
Question: "What happened to the application?"
Answer: "It appears the application was temporarily down because it encountered high traffic and some network issues."
>>>>>>> attempted-AI-integration

Context:
{context}

---
<<<<<<< HEAD

Based on the above context, answer the question: {question}

Please include the source references in your response.
"""

def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("query_text", type=str, nargs='?', default=None, help="The query text.")
    args = parser.parse_args()
    if args.query_text:
        # If a query is provided as an argument, process it once
        answer = query_rag(args.query_text)
        print(answer)
    else:
        # Otherwise, enter live mode
        live_mode()
=======
Based on the above context, answer the question: {question}

Important:
- If no relevant info is found, respond EXACTLY with NO_ANSWER_FOUND.
- Otherwise, provide an answer along with source references.
"""

def hybrid_search(query_text: str, vector_db, k=5):
    """
    Performs a hybrid search combining semantic and (if available) keyword search.
    """
    # Semantic search:
    semantic_results = vector_db.similarity_search_with_score(query_text, k=k)
    
    # Attempt keyword search; if unavailable, fall back to semantic only.
    try:
        keyword_results = vector_db.keyword_search_with_score(query_text, k=k)
    except Exception as e:
        print("Keyword search not available, falling back to semantic search only.")
        keyword_results = []
    
    # Combine results with weighted scores.
    combined_dict = {}
    for doc, score in semantic_results:
        key = id(doc)
        combined_dict[key] = (doc, 0.7 * score)
    for doc, score in keyword_results:
        key = id(doc)
        if key in combined_dict:
            existing_doc, existing_score = combined_dict[key]
            combined_dict[key] = (existing_doc, existing_score + 0.3 * score)
        else:
            combined_dict[key] = (doc, 0.3 * score)
    
    # Sort and return top-k results.
    combined_results = sorted(combined_dict.values(), key=lambda x: x[1], reverse=True)[:k]
    return combined_results

def query_rag(query_text: str, vector_db=None, allow_escalation: bool = True):
    """
    Searches the vector store for relevant context, builds a prompt,
    and invokes the LLM. If the answer is NO_ANSWER_FOUND and escalation is allowed,
    it offers to escalate the query to an SME.
    
    Parameters:
      query_text (str): The user's query.
      vector_db (Chroma, optional): A pre-instantiated vector store. If None, one is created using CHROMA_PATH.
      allow_escalation (bool): Whether to offer escalation when no answer is found.
    
    Returns:
      str: The answer (or escalation result) from the model.
    """
    try:
        embedding_function = get_embedding_function()
        if vector_db is None:
            vector_db = Chroma(
                persist_directory=CHROMA_PATH,
                embedding_function=embedding_function,
            )
        
        # Perform hybrid search on the vector database.
        results = hybrid_search(query_text, vector_db, k=5)
        
        # Build context from the retrieved document chunks.
        context_text = "\n\n---\n\n".join(
            [f"Source {i+1}: {doc.page_content}" for i, (doc, _score) in enumerate(results)]
        )
        
        # Build the prompt.
        prompt_template = ChatPromptTemplate.from_template(PROMPT_TEMPLATE)
        prompt = prompt_template.format(context=context_text, question=query_text)
        
        # Invoke the LLM.
        model = OllamaLLM(model="Mistral")
        response_text = model.invoke(prompt).strip()
        if "no_answer_found" in response_text.lower():
            response_text = "NO_ANSWER_FOUND"
        
        # Gather source references.
        sources = []
        for i, (doc, _score) in enumerate(results):
            source_value = doc.metadata.get("source", f"Source {i+1}")
            if source_value not in sources:
                sources.append(source_value)
        formatted_sources = "\n".join(sources)
        final_response = f"{response_text}\nSources used:\n{formatted_sources}" if sources else response_text
        
        # If no answer was found, and escalation is allowed, offer escalation.
        if final_response.strip() == "NO_ANSWER_FOUND" and allow_escalation:
            choice = input("No info found in context. Escalate to SME? (yes/no): ").lower().strip()
            if choice in ["yes", "y"]:
                sme_answer = escalate_to_sme(query_text)
                if sme_answer:
                    return f"RESPONSE FROM SME: {sme_answer}\n(Stored for future queries)"
                else:
                    return "No SME response received. Please try again later."
        
        return final_response

    except Exception as e:
        return f"An error occurred: {e}"
>>>>>>> attempted-AI-integration

def live_mode():
    print("Entering live mode. Type 'exit' or 'quit' to stop.")
    while True:
        query_text = input("Enter your question: ")
        if query_text.lower() in ['exit', 'quit']:
            print("Exiting...")
            break
<<<<<<< HEAD
        answer = query_rag(query_text)
        print(answer)

def query_rag(query_text: str):
    """
    Returns a string containing the answer from the RAG pipeline 
    plus any source references. 
    """
    try:
        # 1) Prepare the DB
        embedding_function = get_embedding_function()
        db = Chroma(persist_directory=CHROMA_PATH, embedding_function=embedding_function)

        # 2) Search the DB for top-k relevant chunks
        results = db.similarity_search_with_score(query_text, k=5)

        # 3) Build the context from the top-k chunks
        context_text = "\n\n---\n\n".join(
            [f"Source {i + 1}: {doc.page_content}" for i, (doc, _score) in enumerate(results)]
        )

        # 4) Prepare the prompt
        prompt_template = ChatPromptTemplate.from_template(PROMPT_TEMPLATE)
        prompt = prompt_template.format(context=context_text, question=query_text)

        # 5) Use Ollama to generate a response
        model = OllamaLLM(model="Mistral")
        response_text = model.invoke(prompt)

        # 6) Extract the "sources used" from the metadata
        sources = []
        for i, (doc, _score) in enumerate(results):
            source_value = doc.metadata.get("source", f"Source {i + 1}")
            if source_value not in sources:
                sources.append(source_value)

        formatted_sources = "\n".join(sources)
        formatted_response = f" {response_text}\nSources used:\n{formatted_sources}"

        # 7) Return the final response (NOT print)
        return formatted_response

    except Exception as e:
        # Return error message so it shows up in the chatbot
        return f"An error occurred: {e}"
=======
        answer = query_rag(query_text, allow_escalation=True)
        print(answer)

def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("query_text", type=str, nargs="?", default=None, help="The query text.")
    args = parser.parse_args()
    
    if args.query_text:
        answer = query_rag(args.query_text, allow_escalation=True)
        print(answer)
    else:
        live_mode()
>>>>>>> attempted-AI-integration

if __name__ == "__main__":
    main()
